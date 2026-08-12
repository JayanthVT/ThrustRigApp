"""
docx_engine.py — Thrust Test Rig DOCX Report Generator

Structural/visual mirror of pdf_engine.py: same section order (Header →
Meta strip → Initial Parameters → Test Parameter Check → Results →
Observations → Test Charts), same colours, same column-width ratios, same
value-formatting rules (pv/fmt), built from the identical `data` dict and
`chart_images` list the PDF uses — so the two will always show the same
numbers even if Word's rendering engine lays out text slightly differently
than reportlab's. Word/LibreOffice don't give pixel-parity with a PDF
renderer (different font metrics, no fixed page-break control at this
level of the API), but every section, colour, table structure and value
matches. The whole point of shipping this alongside the PDF is that it's
editable — open it, fix a typo, done.
"""

import io
from pathlib import Path

from docx import Document
from docx.shared import Mm, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

BASE_DIR = Path(__file__).parent
LOGO_PATH = BASE_DIR / "assets" / "ideaforge-logo.jpeg"

# ── Colours — same hex values as pdf_engine.py's NAVY/BLUE/MGRAY/LGRAY ──
NAVY = "1B5E20"
BLUE = "2E7D32"
MGRAY = "CCCCCC"
LGRAY = "F5F5F5"
WHITE = "FFFFFF"
BLACK = "000000"
GREEN_TEXT = "1B5E20"
RED_TEXT = "B71C1C"
SUBTITLE_GREEN = "A5D6A7"

FONT = "Calibri"  # closest cross-platform match to reportlab's Helvetica


# ── Low-level oxml helpers (python-docx has no direct cell-shading API) ──
def _shade_cell(cell, hex_color: str):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _cell_border(cell, color=MGRAY, sz=4):
    tcPr = cell._tc.get_or_add_tcPr()
    borders = OxmlElement("w:tcBorders")
    for edge in ("top", "left", "bottom", "right"):
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:color"), color)
        borders.append(el)
    tcPr.append(borders)


def _grid(table, color=MGRAY, sz=4):
    for row in table.rows:
        for cell in row.cells:
            _cell_border(cell, color, sz)


def _set_col_widths(table, widths_mm):
    table.autofit = False
    for row in table.rows:
        for cell, w in zip(row.cells, widths_mm):
            cell.width = Mm(w)


def _cell_text(cell, text, bold=False, italic=False, size=8, color=None,
                align=None, font=FONT, vcenter=True):
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    run = p.add_run(str(text))
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    run.font.name = font
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    if vcenter:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    return run


def _cell_richtext(cell, parts, size=8, align=None, vcenter=True):
    """parts: [(text, bold, color_or_None), ...] — multiple runs in one cell."""
    cell.text = ""
    p = cell.paragraphs[0]
    if align is not None:
        p.alignment = align
    for text, bold, color in parts:
        run = p.add_run(str(text))
        run.bold = bold
        run.font.size = Pt(size)
        run.font.name = FONT
        if color:
            run.font.color.rgb = RGBColor.from_string(color)
    if vcenter:
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER


def _banner(doc, text, bg_hex, w_mm, size=9):
    t = doc.add_table(rows=1, cols=1)
    _set_col_widths(t, [w_mm])
    cell = t.rows[0].cells[0]
    _shade_cell(cell, bg_hex)
    cell.paragraphs[0].paragraph_format.space_before = Pt(2)
    cell.paragraphs[0].paragraph_format.space_after = Pt(2)
    _cell_text(cell, text, bold=True, size=size, color=WHITE)
    return t


def _spacer(doc, pt=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(pt)
    p.paragraph_format.space_before = Pt(0)
    run = p.add_run("")
    run.font.size = Pt(1)


def build_docx_report(data, chart_images, run_name, sections=None) -> bytes:
    """
    Build and return DOCX bytes. Same signature as pdf_engine.build_pdf_report()
    — pass it the exact same data/chart_images/run_name/sections and the two
    documents will contain identical content.
    """
    def _on(key):
        return True if sections is None else sections.get(key, True)

    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Mm(210)
    sec.page_height = Mm(297)
    sec.left_margin = sec.right_margin = Mm(12)
    sec.top_margin = sec.bottom_margin = Mm(12)
    W = 210 - 24  # mm, matches pdf_engine's W = A4[0] - 24mm

    style = doc.styles["Normal"]
    style.font.name = FONT
    style.font.size = Pt(8)

    _BLANK_VALS = {"", "—", "-", "none", "nan", "0", "0.0", "0.00", "0.000"}

    def fmt(key):
        v = data.get(key, "")
        if str(v).strip().lower() in _BLANK_VALS:
            return None
        try:
            fv = float(v)
            if key in ("max_rpm", "target_rpm"):
                v = f"{int(fv):,}"
            elif key == "overall_efficiency":
                v = f"{fv:.4f}"
            elif key == "mechanical_efficiency":
                v = f"{fv:.2f}"
            elif key in ("duration_s", "time_at_target_rpm"):
                v = f"{fv:.1f}"
            else:
                v = f"{fv:.3f}".rstrip("0").rstrip(".")
        except Exception:
            pass
        return str(v)

    # ── HEADER ──
    hdr = doc.add_table(rows=1, cols=3)
    _set_col_widths(hdr, [20, W * 0.62, W - 20 - W * 0.62])
    logo_cell, title_cell, sub_cell = hdr.rows[0].cells
    _shade_cell(logo_cell, WHITE)
    _shade_cell(title_cell, NAVY)
    _shade_cell(sub_cell, NAVY)
    if LOGO_PATH.exists():
        logo_cell.text = ""
        run = logo_cell.paragraphs[0].add_run()
        run.add_picture(str(LOGO_PATH), width=Mm(16), height=Mm(16))
    else:
        _cell_text(logo_cell, "")
    _cell_text(title_cell, run_name, bold=True, size=14, color=WHITE,
               align=WD_ALIGN_PARAGRAPH.CENTER)
    _cell_text(sub_cell, f"Thrust Test Report  ·  {data.get('test_date', '')}",
               size=8, color=SUBTITLE_GREEN, align=WD_ALIGN_PARAGRAPH.LEFT)
    _spacer(doc, 3)

    # ── META STRIP ──
    meta = doc.add_table(rows=1, cols=3)
    _set_col_widths(meta, [W * 0.45, W * 0.20, W * 0.35])
    for cell in meta.rows[0].cells:
        _shade_cell(cell, LGRAY)
    _cell_richtext(meta.rows[0].cells[0],
                   [("File: ", True, None), (data.get("filename", ""), False, None)])
    _cell_richtext(meta.rows[0].cells[1],
                   [("Date: ", True, None), (data.get("test_date", ""), False, None)])
    _cell_richtext(meta.rows[0].cells[2], [
        ("Duration: ", True, None), (f"{data.get('duration_s', '')}s  |  ", False, None),
        ("Rows: ", True, None), (str(data.get("num_rows", "")), False, None),
    ])
    _spacer(doc, 6)

    # ── INITIAL PARAMETERS ──
    if _on("initial_params"):
        ip_groups = [
            ("Reservoir", [
                ("Capacity", "res_capacity", "L"),
                ("Composition", "res_composition", ""),
                ("Temperature", "res_temperature", "°C"),
            ]),
            ("Duty Cycle", [("Duty Cycle & Flowrate", "duty_cycle", "")]),
            ("Temperature", [
                ("Initial ESC Temp", "init_esc_temp", "°C"),
                ("Initial Motor Temp", "init_motor_temp", "°C"),
                ("Ambient", "ambient_temp", "°C"),
                ("ESC Inlet Coolant", "esc_inlet_coolant", "°C"),
                ("Motor Inlet Coolant", "motor_inlet_coolant", "°C"),
            ]),
            ("Flowrate", [
                ("ESC Inlet", "esc_inlet_flow", "LPM"),
                ("Motor Inlet", "motor_inlet_flow", "LPM"),
            ]),
            ("Pressure", [
                ("ESC Inlet", "esc_inlet_pressure", "Bar"),
                ("Motor Inlet", "motor_inlet_pressure", "Bar"),
            ]),
            ("Battery", [
                ("Battery Voltage", "battery_voltage", "V"),
                ("SOC", "battery_soc", ""),
                ("SOH", "battery_soh", ""),
            ]),
            ("Fintube", [
                ("Inlet Temperature", "fin_inlet_temp", "°C"),
                ("Outlet Temperature", "fin_outlet_temp", "°C"),
            ]),
        ]
        total_rows = sum(len(rows) for _, rows in ip_groups)

        _banner(doc, "INITIAL PARAMETERS", BLUE, W)

        t_ip = doc.add_table(rows=total_rows, cols=4)
        _set_col_widths(t_ip, [W * 0.18, W * 0.32, W * 0.38, W * 0.12])
        _grid(t_ip)

        r = 0
        for grp_label, rows in ip_groups:
            span_start = r
            for i, (lbl, key, unit) in enumerate(rows):
                val = fmt(key)
                row_cells = t_ip.rows[r].cells
                _shade_cell(row_cells[0], LGRAY)
                _cell_text(row_cells[0], grp_label if i == 0 else "",
                           bold=True, size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
                _cell_text(row_cells[1], lbl, bold=True, size=8)
                _cell_text(row_cells[2], val or "", size=8)
                _cell_text(row_cells[3], unit if val else "", size=7, color="333333")
                r += 1
            if len(rows) > 1:
                t_ip.cell(span_start, 0).merge(t_ip.cell(r - 1, 0))

        _spacer(doc, 6)

    # ── TEST PARAMETER CHECK ──
    tpc_rows = data.get("test_param_check")
    if isinstance(tpc_rows, str):
        try:
            import json as _json
            tpc_rows = _json.loads(tpc_rows)
        except Exception:
            tpc_rows = []

    if tpc_rows and _on("test_param_check"):
        _banner(doc, "TEST PARAMETER CHECK", BLUE, W)
        t_tpc = doc.add_table(rows=1 + len(tpc_rows), cols=3)
        _set_col_widths(t_tpc, [W * 0.55, W * 0.12, W * 0.33])
        _grid(t_tpc)

        hdr_cells = t_tpc.rows[0].cells
        _shade_cell(hdr_cells[0], LGRAY)
        _shade_cell(hdr_cells[1], LGRAY)
        _shade_cell(hdr_cells[2], LGRAY)
        _cell_text(hdr_cells[0], "Parameter / Criteria", bold=True, size=8)
        _cell_text(hdr_cells[1], "Pass / Fail", bold=True, size=8)
        _cell_text(hdr_cells[2], "Remarks", bold=True, size=8)

        for i, row in enumerate(tpc_rows, start=1):
            passed = row.get("passed", False)
            failed = row.get("failed", False)
            pf_text = "Pass" if passed else ("Fail" if failed else "—")
            pf_color = GREEN_TEXT if passed else (RED_TEXT if failed else None)
            remarks = str(row.get("remarks", ""))
            if passed and not remarks:
                remarks = "OKAY"
            cells = t_tpc.rows[i].cells
            _cell_text(cells[0], str(row.get("criteria", "")), size=8)
            _cell_text(cells[1], pf_text, bold=True, size=8, color=pf_color)
            _cell_text(cells[2], remarks, size=8)
        _spacer(doc, 6)

    # ── RESULTS (core) + MEASURABLE PARAMS (efficiency) ──
    res_rows_core = [
        ("Max. Temp — ESC Inlet", "max_esc_inlet_temp", "°C"),
        ("Max. Temp — Motor Inlet", "max_motor_inlet_temp", "°C"),
        ("Max. Pressure — ESC Inlet", "max_esc_pressure", "Bar"),
        ("Battery Voltage (post-run)", "battery_voltage_post", "V"),
        ("Max. RPM", "max_rpm", "RPM"),
        ("Max. Torque", "max_torque", "Nm"),
        ("Max. Thrust", "max_thrust", "N"),
        ("Fin Tube Inlet Temp (max)", "max_fin_inlet_temp", "°C"),
        ("Fin Tube Outlet Temp (max)", "max_fin_outlet_temp", "°C"),
        ("Max. ESC Temp", "max_esc_temp", "°C"),
        ("Max. Motor Temp", "max_motor_temp", "°C"),
        ("Time at Target RPM", "time_at_target_rpm", "s"),
    ]
    res_rows_measurable = [
        ("Mechanical Power", "mechanical_power", "W"),
        ("Electrical Power", "electrical_power", "W"),
        ("Mechanical Efficiency", "mechanical_efficiency", "%"),
        ("Overall Efficiency", "overall_efficiency", "g/W"),
    ]

    def _rows_for(defs):
        out = []
        for lbl, key, unit in defs:
            v = fmt(key)
            if v is None:
                continue
            out.append((lbl, v, unit))
        return out

    res_data = []
    if _on("results"):
        res_data += _rows_for(res_rows_core)
    if _on("measurable_params"):
        res_data += _rows_for(res_rows_measurable)

    if _on("results") or _on("measurable_params"):
        _banner(doc, "RESULTS", NAVY, W)
        if res_data:
            t_res = doc.add_table(rows=len(res_data), cols=3)
            _set_col_widths(t_res, [W * 0.55, W * 0.33, W * 0.12])
            _grid(t_res)
            for i, (lbl, val, unit) in enumerate(res_data):
                cells = t_res.rows[i].cells
                _shade_cell(cells[0], LGRAY)
                _cell_text(cells[0], lbl, bold=True, size=8, color=NAVY)
                _cell_text(cells[1], val, size=8)
                _cell_text(cells[2], unit, size=7, color="333333")
        else:
            doc.add_paragraph("No results recorded for this run.")
        _spacer(doc, 6)

    # ── OBSERVATIONS ──
    obs = data.get("notes", "")
    if obs and _on("observations"):
        _banner(doc, "OBSERVATIONS", BLUE, W)
        t_obs = doc.add_table(rows=1, cols=1)
        _set_col_widths(t_obs, [W])
        _grid(t_obs)
        cell = t_obs.rows[0].cells[0]
        cell.text = ""
        for j, line in enumerate(str(obs).replace("\r\n", "\n").replace("\r", "\n").split("\n")):
            p = cell.paragraphs[0] if j == 0 else cell.add_paragraph()
            run = p.add_run(line)
            run.font.size = Pt(8)
            run.font.name = FONT
        _spacer(doc, 6)

    # ── CHARTS ──
    if chart_images and _on("charts"):
        _banner(doc, "TEST CHARTS", NAVY, W)
        _spacer(doc, 4)
        for chart_title, png_bytes in chart_images:
            p = doc.add_paragraph()
            run = p.add_run(chart_title)
            run.bold = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor.from_string(NAVY)
            img_p = doc.add_paragraph()
            img_run = img_p.add_run()
            img_run.add_picture(io.BytesIO(png_bytes), width=Mm(W))
            _spacer(doc, 4)

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.read()
